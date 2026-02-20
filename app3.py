import streamlit as st
import pandas as pd
from datetime import datetime, timedelta, timezone
import gspread
from google.oauth2.service_account import Credentials
import io
import os

st.set_page_config(page_title="Sistema de Gesti√≥n de RH DFC", page_icon="üìÖ", layout="wide")

SCOPES = ['https://www.googleapis.com/auth/spreadsheets', 'https://www.googleapis.com/auth/drive']

NORMATIVA = {
    'economico': {
        'nombre': 'D√≠a Econ√≥mico', 
        'max_dias': 3, 
        'max_ocasiones': 3, 
        'intervalo_dias': 30, 
        'descripcion': 'Hasta 3 ocasiones por a√±o',
        'limite': '3 ocasiones/a√±o'
    },
    'matrimonio': {
        'nombre': 'Matrimonio', 
        'max_dias': 10, 
        'max_ocasiones': 1, 
        'descripcion': 'Por una sola ocasi√≥n en la vida',
        'limite': '1 vez en la vida'
    },
    'fallecimiento': {
        'nombre': 'Fallecimiento/Enfermedad Grave', 
        'max_dias': 5, 
        'descripcion': 'Parientes primer grado',
        'limite': 'Sin l√≠mite'
    },
    'jubilacion': {
        'nombre': 'Tr√°mites Jubilaci√≥n', 
        'max_dias': 2, 
        'descripcion': 'Solo cuando se jubila',
        'limite': '1 vez en la vida'
    },
    'examen': {
        'nombre': 'Examen Profesional/Tesis', 
        'max_dias': 3, 
        'descripcion': 'Presentaci√≥n de grado',
        'limite': 'M√°ximo 3 veces'
    },
    'mudanza': {
        'nombre': 'Cambio de Domicilio', 
        'max_dias': 1, 
        'descripcion': 'Para mudanza',
        'limite': '2 veces/a√±o'
    }
}

def conectar_sheets():
    try:
        creds = Credentials.from_service_account_info(st.secrets["google_sheets"], scopes=SCOPES)
        return gspread.authorize(creds)
    except:
        return None

def verificar_login(usuario, password):
    try:
        usuarios = st.secrets["usuarios"]
        if usuario in usuarios and usuarios[usuario]["password"] == password:
            tipo = usuarios[usuario].get("tipo", "admin")  # Default admin
            return True, usuarios[usuario]["nombre"], tipo
    except:
        pass
    return False, None, None

def inicializar_sheets(client):
    try:
        spreadsheet = client.open("Dias_Economicos_Formacion_Continua")
        sheet_empleados = spreadsheet.worksheet("Empleados")
        sheet_solicitudes = spreadsheet.worksheet("Solicitudes")
        
        # Hoja de Incapacidades
        try:
            sheet_incapacidades = spreadsheet.worksheet("Incapacidades")
        except gspread.WorksheetNotFound:
            sheet_incapacidades = spreadsheet.add_worksheet(title="Incapacidades", rows=1000, cols=20)
            sheet_incapacidades.update('A1:S1', [[
                'ID', 'EmpleadoID', 'RFC', 'Nombre Completo', 'Correo Empleado', 'Telefono Contacto',
                'Numero Incapacidad', 'Fecha Inicio', 'Fecha Termino', 'Dias Totales',
                'Tipo Incapacidad', 'Excede Dias', 'Dias Enfermedad General', 'Dias Maternidad',
                'Dias Riesgo Trabajo', 'Dias Posible Riesgo', 'Mes Correspondiente', 
                'Estado', 'Registrado Por'
            ]])
        
        # Nueva hoja de Pendientes por Empleado
        try:
            sheet_pendientes = spreadsheet.worksheet("Pendientes_Empleado")
        except gspread.WorksheetNotFound:
            sheet_pendientes = spreadsheet.add_worksheet(title="Pendientes_Empleado", rows=500, cols=12)
            sheet_pendientes.update('A1:L1', [[
                'ID', 'EmpleadoID', 'RFC', 'Nombre Completo', 'Tipo_Pendiente', 
                'Descripcion', 'Quincena', 'A√±o', 'Estado', 'Fecha_Registro', 
                'Fecha_Completado', 'Completado_Por'
            ]])

        try:
            sheet_constancias = spreadsheet.worksheet("Constancias")
        except gspread.WorksheetNotFound:
            sheet_constancias = spreadsheet.add_worksheet(title="Constancias", rows=100, cols=20)
            sheet_constancias.update('A1:Q1', [[
                'Hoja', 'Nombre Completo', 'Apellido paterno', 'Apellido Materno', 'Nombre(s)',
                'N.C.T. Adscripci√≥n', 'C.C.T. ADSCRIPCI√ìN', 'Clave Presupuestal', 'RFC',
                'INGRESOA LA SEJ', 'Nombramiento', 'Descripci√≥n de puesto', 
                'Se desempe√±a en', 'Subsitema', 'HORARIO', 'TEL. PERSONAL', 'TEL. ext.'
            ]])

        return spreadsheet, sheet_empleados, sheet_solicitudes, sheet_incapacidades, sheet_pendientes, sheet_constancias
    except Exception as e:
        st.error(f"ERROR en inicializar_sheets: {str(e)}")
        import traceback
        st.code(traceback.format_exc())
        return None, None, None, None, None, None

def cargar_datos_con_calculo(sheet_emp, sheet_sol):
    """Carga datos y CALCULA d√≠as disponibles en tiempo real"""
    df_emp = pd.DataFrame(sheet_emp.get_all_records())
    df_sol = pd.DataFrame(sheet_sol.get_all_records())
    
    # CALCULAR D√çAS DISPONIBLES REALES
    for idx, emp in df_emp.iterrows():
        emp_id = emp['ID']
        
        # Contar d√≠as econ√≥micos usados este a√±o
        solicitudes_emp = df_sol[
            (df_sol['EmpleadoID'] == emp_id) & 
            (df_sol['Tipo Permiso'] == 'economico')
        ]
        
        dias_usados = 0
        if len(solicitudes_emp) > 0:
            solicitudes_emp['Fecha_Reg'] = pd.to_datetime(solicitudes_emp['Fecha Registro'], errors='coerce')
            solicitudes_a√±o = solicitudes_emp[solicitudes_emp['Fecha_Reg'].dt.year == datetime.now().year]
            dias_usados = int(solicitudes_a√±o['Dias Solicitados'].sum())
        
        df_emp.at[idx, 'DIAS_REALES'] = 9 - dias_usados
    
    return df_emp, df_sol

def validar_solicitud(emp_id, tipo, dias, fecha_inicio, df_emp, df_sol):
    """Validaci√≥n completa de solicitud"""
    errores = []
    advertencias = []
    
    emp_info = df_emp[df_emp['ID'] == emp_id].iloc[0]
    dias_disponibles = int(emp_info['DIAS_REALES'])
    config = NORMATIVA[tipo]
    
    # Validar d√≠as m√°ximos
    if dias > config['max_dias']:
        errores.append(f"‚ùå M√°ximo permitido: {config['max_dias']} d√≠as")
    
    if tipo == 'economico':
        # Validar d√≠as disponibles
        if dias > dias_disponibles:
            errores.append(f"‚ùå Solo tiene {dias_disponibles} d√≠as disponibles (solicit√≥ {dias})")
        
        # Validar ocasiones en el a√±o
        a√±o_actual = datetime.now().year
        solicitudes_eco = df_sol[
            (df_sol['EmpleadoID'] == emp_id) &
            (df_sol['Tipo Permiso'] == 'economico')
        ]
        
        if len(solicitudes_eco) > 0:
            solicitudes_eco['Fecha_Reg'] = pd.to_datetime(solicitudes_eco['Fecha Registro'], errors='coerce')
            solicitudes_a√±o = solicitudes_eco[solicitudes_eco['Fecha_Reg'].dt.year == a√±o_actual]
            
            if len(solicitudes_a√±o) >= config['max_ocasiones']:
                errores.append(f"‚ùå Ya alcanz√≥ el l√≠mite de {config['max_ocasiones']} ocasiones en el a√±o")
            
            # Validar intervalo 30 d√≠as
            if len(solicitudes_eco) > 0:
                solicitudes_eco['Fecha_Fin'] = pd.to_datetime(solicitudes_eco['Fecha Fin'], errors='coerce')
                ultima_fecha_fin = solicitudes_eco['Fecha_Fin'].max()
                fecha_inicio_dt = pd.to_datetime(fecha_inicio)
                dias_diferencia = (fecha_inicio_dt - ultima_fecha_fin).days
                
                if dias_diferencia < 30:
                    fecha_valida = ultima_fecha_fin + timedelta(days=30)
                    errores.append(
                        f"‚ùå Debe esperar {30 - dias_diferencia} d√≠as m√°s\n"
                        f"   √öltimo d√≠a usado: {ultima_fecha_fin.strftime('%d/%m/%Y')}\n"
                        f"   Puede solicitar desde: {fecha_valida.strftime('%d/%m/%Y')}"
                    )
        
        # Advertencia
        if dias_disponibles - dias <= 2 and dias <= dias_disponibles:
            advertencias.append(f"‚ö†Ô∏è Despu√©s quedar√°n {dias_disponibles - dias} d√≠as disponibles")
    
    # Matrimonio solo una vez EN LA VIDA
    if tipo == 'matrimonio':
        solicitudes_mat = df_sol[
            (df_sol['EmpleadoID'] == emp_id) &
            (df_sol['Tipo Permiso'] == 'matrimonio')
        ]
        if len(solicitudes_mat) > 0:
            errores.append("‚ùå La licencia por matrimonio solo se otorga UNA VEZ en la vida")
    
    # Jubilaci√≥n solo una vez EN LA VIDA
    if tipo == 'jubilacion':
        solicitudes_jub = df_sol[
            (df_sol['EmpleadoID'] == emp_id) &
            (df_sol['Tipo Permiso'] == 'jubilacion')
        ]
        if len(solicitudes_jub) > 0:
            errores.append("‚ùå La licencia por jubilaci√≥n solo se otorga UNA VEZ (cuando se jubila)")
    
    # Examen profesional: m√°ximo 3 veces en la vida (licenciatura, maestr√≠a, doctorado)
    if tipo == 'examen':
        solicitudes_exam = df_sol[
            (df_sol['EmpleadoID'] == emp_id) &
            (df_sol['Tipo Permiso'] == 'examen')
        ]
        if len(solicitudes_exam) >= 3:
            errores.append("‚ùå La licencia por examen profesional se otorga m√°ximo 3 veces (licenciatura, maestr√≠a, doctorado)")
    
    # Mudanza: m√°ximo 2 veces por a√±o (razonable)
    if tipo == 'mudanza':
        a√±o_actual = datetime.now().year
        solicitudes_mud = df_sol[
            (df_sol['EmpleadoID'] == emp_id) &
            (df_sol['Tipo Permiso'] == 'mudanza')
        ]
        if len(solicitudes_mud) > 0:
            solicitudes_mud['Fecha_Reg'] = pd.to_datetime(solicitudes_mud['Fecha Registro'], errors='coerce')
            solicitudes_a√±o = solicitudes_mud[solicitudes_mud['Fecha_Reg'].dt.year == a√±o_actual]
            if len(solicitudes_a√±o) >= 2:
                errores.append("‚ùå La licencia por mudanza se otorga m√°ximo 2 veces por a√±o")
    
    return errores, advertencias

def generar_reporte_completo_mes(df_emp, df_sol, df_incap, df_pend, mes, a√±o):
    """Genera un Excel completo con TODO el mes"""
    output = io.BytesIO()
    
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        # Filtrar datos del mes
        df_sol_mes = df_sol.copy()
        if len(df_sol_mes) > 0:
            df_sol_mes['Fecha_Reg'] = pd.to_datetime(df_sol_mes['Fecha Registro'], errors='coerce')
            df_sol_mes = df_sol_mes[
                (df_sol_mes['Fecha_Reg'].dt.month == mes) & 
                (df_sol_mes['Fecha_Reg'].dt.year == a√±o)
            ]
        
        df_incap_mes = df_incap.copy()
        if len(df_incap_mes) > 0 and 'Fecha Inicio' in df_incap_mes.columns:
            df_incap_mes['Fecha_Inicio_dt'] = pd.to_datetime(df_incap_mes['Fecha Inicio'], errors='coerce')
            df_incap_mes = df_incap_mes[
                (df_incap_mes['Fecha_Inicio_dt'].dt.month == mes) & 
                (df_incap_mes['Fecha_Inicio_dt'].dt.year == a√±o)
            ]
        
        df_pend_mes = df_pend.copy()
        if len(df_pend_mes) > 0 and 'Fecha_Registro' in df_pend_mes.columns:
            df_pend_mes['Fecha_Reg_dt'] = pd.to_datetime(df_pend_mes['Fecha_Registro'], errors='coerce')
            df_pend_mes = df_pend_mes[
                (df_pend_mes['Fecha_Reg_dt'].dt.month == mes) & 
                (df_pend_mes['Fecha_Reg_dt'].dt.year == a√±o)
            ]
        
        # HOJA 1: RESUMEN EJECUTIVO
        resumen_data = {
            'INDICADOR': [
                'Total Solicitudes del Mes',
                'Total Incapacidades del Mes',
                'Total Pendientes Registrados',
                'D√≠as Econ√≥micos Solicitados',
                'Otros Permisos Solicitados',
                'Total D√≠as Solicitados',
                'Empleados que Solicitaron',
                'Pendientes Activos'
            ],
            'VALOR': [
                len(df_sol_mes),
                len(df_incap_mes),
                len(df_pend_mes),
                len(df_sol_mes[df_sol_mes['Tipo Permiso'] == 'economico']) if len(df_sol_mes) > 0 else 0,
                len(df_sol_mes[df_sol_mes['Tipo Permiso'] != 'economico']) if len(df_sol_mes) > 0 else 0,
                df_sol_mes['Dias Solicitados'].sum() if len(df_sol_mes) > 0 else 0,
                df_sol_mes['EmpleadoID'].nunique() if len(df_sol_mes) > 0 else 0,
                len(df_pend_mes[df_pend_mes['Estado'] == 'Pendiente']) if len(df_pend_mes) > 0 else 0
            ]
        }
        df_resumen = pd.DataFrame(resumen_data)
        df_resumen.to_excel(writer, sheet_name='RESUMEN', index=False)
        
        # HOJA 2: SOLICITUDES DEL MES
        if len(df_sol_mes) > 0:
            df_sol_export = df_sol_mes[[
                'ID', 'EmpleadoID', 'RFC', 'Nombre Completo',
                'Tipo Permiso', 'Fecha Inicio', 'Fecha Fin', 'Dias Solicitados',
                'Motivo', 'Fecha Registro', 'Aprobado Por', 'Registrado Por'
            ]].copy()
            df_sol_export.to_excel(writer, sheet_name='Solicitudes', index=False)
        
        # HOJA 3: INCAPACIDADES DEL MES
        if len(df_incap_mes) > 0:
            df_incap_export = df_incap_mes.drop(columns=['Fecha_Inicio_dt'], errors='ignore')
            df_incap_export.to_excel(writer, sheet_name='Incapacidades', index=False)
        
        # HOJA 4: PENDIENTES DEL MES
        if len(df_pend_mes) > 0:
            df_pend_export = df_pend_mes.drop(columns=['Fecha_Reg_dt'], errors='ignore')
            df_pend_export.to_excel(writer, sheet_name='Pendientes', index=False)
        
        # HOJA 5: ESTAD√çSTICAS POR TIPO
        if len(df_sol_mes) > 0:
            stats_tipo = df_sol_mes.groupby('Tipo Permiso').agg({
                'ID': 'count',
                'Dias Solicitados': 'sum',
                'EmpleadoID': 'nunique'
            }).rename(columns={
                'ID': 'Num Solicitudes',
                'Dias Solicitados': 'Total Dias',
                'EmpleadoID': 'Num Empleados'
            })
            stats_tipo.to_excel(writer, sheet_name='Stats por Tipo')
        
        # HOJA 6: ESTAD√çSTICAS POR EMPLEADO
        if len(df_sol_mes) > 0:
            stats_emp = df_sol_mes.groupby(['EmpleadoID', 'Nombre Completo']).agg({
                'ID': 'count',
                'Dias Solicitados': 'sum'
            }).rename(columns={
                'ID': 'Num Solicitudes',
                'Dias Solicitados': 'Total Dias'
            })
            stats_emp.to_excel(writer, sheet_name='Stats por Empleado')
        
        # HOJA 7: ESTADO ACTUAL DE EMPLEADOS
        df_emp_export = df_emp[['ID', 'RFC', 'PATERNO', 'MATERNO', 'NOMBRE', 'CURP', 'PLAZA', 'DIAS_REALES']].copy()
        df_emp_export.to_excel(writer, sheet_name='Estado Empleados', index=False)
    
    output.seek(0)
    return output.getvalue()

def crear_trazabilidad_completa(df_sol, df_emp):
    """Crea un reporte de trazabilidad completo"""
    if len(df_sol) == 0:
        return pd.DataFrame()
    
    df_traz = df_sol.copy()
    df_traz['Fecha_Reg'] = pd.to_datetime(df_traz['Fecha Registro'], errors='coerce')
    
    # Enriquecer con informaci√≥n del empleado
    df_traz = df_traz.merge(
        df_emp[['ID', 'PLAZA']],
        left_on='EmpleadoID',
        right_on='ID',
        how='left',
        suffixes=('', '_emp')
    )
    
    # Ordenar cronol√≥gicamente
    df_traz = df_traz.sort_values('Fecha_Reg', ascending=False)
    
    # Columnas de trazabilidad
    columnas_traz = [
        'Fecha Registro', 'Nombre Completo', 'RFC', 'PLAZA',
        'Tipo Permiso', 'Fecha Inicio', 'Fecha Fin', 'Dias Solicitados',
        'Motivo', 'Aprobado Por', 'Registrado Por'
    ]
    
    return df_traz[columnas_traz]

def generar_alertas(df_empleados):
    """Genera alertas de empleados con pocos d√≠as"""
    alertas = []
    for _, emp in df_empleados.iterrows():
        dias = int(emp['DIAS_REALES'])
        nombre = f"{emp['PATERNO']} {emp['MATERNO']} {emp['NOMBRE']}"
        
        if dias == 0:
            alertas.append({'tipo': 'error', 'mensaje': f"üö´ {nombre} NO tiene d√≠as disponibles"})
        elif dias == 1:
            alertas.append({'tipo': 'warning', 'mensaje': f"‚ö†Ô∏è {nombre} tiene solo 1 d√≠a disponible"})
        elif dias <= 3:
            alertas.append({'tipo': 'info', 'mensaje': f"‚ÑπÔ∏è {nombre} tiene {dias} d√≠as disponibles"})
    
    return alertas

def verificar_fechas_limite():
    """Recordatorios de fechas l√≠mite para propuestas"""
    zona_mexico = timezone(timedelta(hours=-6))
    hoy = datetime.now(zona_mexico).date()
    
    # CALENDARIO ESTATAL
    fechas_estatal = {
        'Q03': datetime(2026, 1, 23).date(),
        'Q04': datetime(2026, 2, 5).date(),
        'Q05': datetime(2026, 2, 16).date(),
        'Q06': datetime(2026, 3, 3).date(),
        'Q07': datetime(2026, 3, 3).date(),
        'Q08': datetime(2026, 3, 23).date(),
        'Q09': datetime(2026, 4, 20).date(),
        'Q10': datetime(2026, 5, 7).date(),
        'Q11': datetime(2026, 5, 20).date(),
        'Q12': datetime(2026, 6, 8).date(),
        'Q13': datetime(2026, 6, 22).date(),
        'Q14': datetime(2026, 7, 8).date(),
        'Q15': datetime(2026, 7, 21).date(),
        'Q16': datetime(2026, 8, 10).date(),
        'Q17': datetime(2026, 8, 21).date(),
        'Q18': datetime(2026, 9, 4).date(),
        'Q19': datetime(2026, 9, 22).date(),
        'Q20': datetime(2026, 10, 6).date(),
        'Q21': datetime(2026, 10, 21).date(),
        'Q22': datetime(2026, 11, 5).date(),
    }
    
    # CALENDARIO FEDERALIZADO
    fechas_federal = {
        'Q02': datetime(2026, 1, 9).date(),
        'Q03': datetime(2026, 1, 23).date(),
        'Q04': datetime(2026, 2, 5).date(),
        'Q05': datetime(2026, 2, 19).date(),
        'Q06': datetime(2026, 3, 6).date(),
        'Q07': datetime(2026, 3, 6).date(),
        'Q08': datetime(2026, 4, 7).date(),
        'Q09': datetime(2026, 4, 23).date(),
        'Q10': datetime(2026, 5, 7).date(),
        'Q11': datetime(2026, 5, 22).date(),
        'Q12': datetime(2026, 6, 8).date(),
        'Q13': datetime(2026, 6, 23).date(),
        'Q14': datetime(2026, 6, 23).date(),
        'Q15': datetime(2026, 6, 23).date(),
        'Q16': datetime(2026, 8, 6).date(),
        'Q17': datetime(2026, 8, 21).date(),
        'Q18': datetime(2026, 9, 7).date(),
        'Q19': datetime(2026, 9, 22).date(),
        'Q20': datetime(2026, 10, 7).date(),
        'Q21': datetime(2026, 10, 22).date(),
        'Q22': datetime(2026, 11, 5).date(),
        'Q23': datetime(2026, 11, 20).date(),
        'Q24': datetime(2026, 11, 20).date(),
    }
    
    alertas = {'criticas': [], 'proximas': [], 'futuras': []}
    
    # ESTATAL
    for qna, fecha in fechas_estatal.items():
        dias = (fecha - hoy).days
        if dias < 0:
            continue
        
        item = {
            'sistema': 'ESTATAL',
            'quincena': qna,
            'fecha': fecha.strftime('%d/%m/%Y'),
            'dias': dias
        }
        
        if 0 <= dias <= 5:
            alertas['criticas'].append(item)
        elif 4 <= dias <= 15:
            alertas['proximas'].append(item)
        elif 16 <= dias <= 90:
            alertas['futuras'].append(item)
    
    # FEDERALIZADO
    for qna, fecha in fechas_federal.items():
        dias = (fecha - hoy).days
        if dias < 0:
            continue
        
        item = {
            'sistema': 'FEDERAL',
            'quincena': qna,
            'fecha': fecha.strftime('%d/%m/%Y'),
            'dias': dias
        }
        
        if 0 <= dias <= 5:
            alertas['criticas'].append(item)
        elif 4 <= dias <= 15:
            alertas['proximas'].append(item)
        elif 16 <= dias <= 90:
            alertas['futuras'].append(item)
    
    return alertas

def generar_constancias_word(df_constancias, empleados_seleccionados, num_quincena, a√±o, fecha_elaboracion):
    """Genera documento Word con constancias conservando formato e im√°genes"""
    from docx import Document
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    import os
    
    plantilla_path = os.path.join(os.path.dirname(__file__), 'templates', 'plantilla.docx')
    
    if not os.path.exists(plantilla_path):
        raise FileNotFoundError(f"No se encontr√≥ la plantilla en: {plantilla_path}")
    
    meses = {
        1: 'enero', 2: 'febrero', 3: 'marzo', 4: 'abril',
        5: 'mayo', 6: 'junio', 7: 'julio', 8: 'agosto',
        9: 'septiembre', 10: 'octubre', 11: 'noviembre', 12: 'diciembre'
    }
    fecha_texto = f"{fecha_elaboracion.day} de {meses[fecha_elaboracion.month]} de {fecha_elaboracion.year}"
    
    # Lista para guardar documentos individuales
    docs = []
    
    for nombre_empleado in empleados_seleccionados:
        # Filtrar TODOS los registros de este empleado
        registros_empleado = df_constancias[df_constancias['Nombre Completo'] == nombre_empleado]
        
        # Generar una constancia por CADA registro
        for idx, emp in registros_empleado.iterrows():
            # Abrir plantilla NUEVA para cada registro
            doc = Document(plantilla_path)
        
        tel_personal = str(emp['TEL. PERSONAL'])
        if '.' in tel_personal:
            try:
                tel_personal = f"{int(float(tel_personal)):010d}"
            except:
                pass
        
        reemplazos = {
            '<<QUINCENA>>': str(num_quincena),
            '<<A√ëO>>': str(a√±o),
            '<<FECHA>>': fecha_texto,
            '<<APELLIDO_PATERNO>>': str(emp['Apellido paterno']),
            '<<APELLIDO_MATERNO>>': str(emp['Apellido Materno']),
            '<<NOMBRE>>': str(emp['Nombre(s)']),
            '<<RFC>>': str(emp['RFC']),
            '<<FECHA_INGRESO>>': str(emp['INGRESOA LA SEJ']),
            '<<SE_DESEMPENA_EN>>': str(emp['Se desempe√±a en']),
            '<<DESCRIPCION_PUESTO>>': str(emp['Descripci√≥n de puesto']),
            '<<CCT>>': str(emp['C.C.T. ADSCRIPCI√ìN']),
            '<<CLAVE_PRESUPUESTAL>>': str(emp['Clave Presupuestal']),
            '<<TEL_PERSONAL>>': tel_personal,
            '<<TEL_EXT>>': str(emp['TEL. ext.']),
            '<<HOJA>>': str(int(emp['Hoja']))
        }
        
        # Reemplazar en p√°rrafos
        for paragraph in doc.paragraphs:
            texto_completo = paragraph.text
            for marcador, valor in reemplazos.items():
                if marcador in texto_completo:
                    texto_completo = texto_completo.replace(marcador, valor)
            
            # Actualizar el texto
            if texto_completo != paragraph.text:
                for run in paragraph.runs:
                    run.text = ''
                if paragraph.runs:
                    paragraph.runs[0].text = texto_completo
                else:
                    paragraph.add_run(texto_completo)
        
        # Reemplazar en tablas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        texto_completo = paragraph.text
                        for marcador, valor in reemplazos.items():
                            if marcador in texto_completo:
                                texto_completo = texto_completo.replace(marcador, valor)
                        
                        if texto_completo != paragraph.text:
                            for run in paragraph.runs:
                                run.text = ''
                            if paragraph.runs:
                                paragraph.runs[0].text = texto_completo
                            else:
                                paragraph.add_run(texto_completo)
        
        docs.append(doc)
    
    # Combinar documentos SIN docxcompose
    doc_final = docs[0]
    
    for doc in docs[1:]:
        # NO agregar page_break - solo copiar contenido
        for element in list(doc.element.body):
            # SALTAR sectPr (propiedades de secci√≥n)
            if element.tag.endswith('sectPr'):
                continue
            doc_final.element.body.append(element)
    
    output_path = os.path.join(os.path.dirname(__file__), f'Constancias_Q{num_quincena}_{a√±o}.docx')
    doc_final.save(output_path)

    return output_path
    

def convertir_word_a_pdf(word_path):
    """Convierte Word a PDF usando LibreOffice directamente"""
    import subprocess
    import os
    
    # Definir ruta de salida
    output_dir = os.path.dirname(os.path.abspath(word_path))
    pdf_path = word_path.replace('.docx', '.pdf')
    
    try:
        # Comando directo de LibreOffice (lowriter es m√°s espec√≠fico para docs)
        comando = [
            'lowriter',
            '--headless',
            '--convert-to', 'pdf',
            '--outdir', output_dir,
            '-env:UserInstallation=file:///tmp/lo_conversion',
            os.path.abspath(word_path)
        ]
        
        # Ejecutar con timeout de 30 segundos
        result = subprocess.run(comando, capture_output=True, text=True, timeout=30)
        
        if os.path.exists(pdf_path):
            return pdf_path
        else:
            st.error(f"Error de LibreOffice: {result.stderr}")
            return None
            
    except Exception as e:
        st.error(f"No se pudo ejecutar la conversi√≥n: {e}")
        return None
    
def generar_comisiones_word(df_comisiones, tipo_comision, oficio_inicial, fecha_doc, fecha_inicio, fecha_fin):
    """Genera documento Word de comisiones"""
    from docx import Document
    from datetime import datetime
    import os
    
    try:
        # Seleccionar plantilla seg√∫n tipo
        if tipo_comision == "Encargados CM":
            plantilla_path = os.path.join(os.path.dirname(__file__), 'templates', 'PLANTILLA_ENCARGADOS_CM.docx')
        else:
            plantilla_path = os.path.join(os.path.dirname(__file__), 'templates', 'PLANTILLA_COMISIONES_GENERALES.docx')
        
        if not os.path.exists(plantilla_path):
            raise FileNotFoundError(f"Plantilla no encontrada: {plantilla_path}")
        
        # Generar un documento por cada persona
        docs = []
        oficio_actual = oficio_inicial
        
        for idx, persona in df_comisiones.iterrows():
            doc = Document(plantilla_path)
            
            # Diccionario de meses en espa√±ol
            meses = {
                1: 'enero', 2: 'febrero', 3: 'marzo', 4: 'abril',
                5: 'mayo', 6: 'junio', 7: 'julio', 8: 'agosto',
                9: 'septiembre', 10: 'octubre', 11: 'noviembre', 12: 'diciembre'
            }
            
            # Convertir fechas a espa√±ol
            fecha_doc_esp = f"{fecha_doc.day} de {meses[fecha_doc.month]} de {fecha_doc.year}"
            fecha_inicio_esp = f"{fecha_inicio.day} de {meses[fecha_inicio.month]} de {fecha_inicio.year}"
            fecha_fin_esp = f"{fecha_fin.day} de {meses[fecha_fin.month]} de {fecha_fin.year}"
            
            # Preparar reemplazos
            reemplazos = {
                '<<OFICIO>>': f"{oficio_actual}/52/2026",
                '<<FECHA>>': fecha_doc_esp,
                '<<NOMBRE_COMPLETO>>': persona['nombre_completo'],
                '<<FECHA_INICIO>>': fecha_inicio_esp,
                '<<FECHA_FIN>>': fecha_fin_esp
            }
            
            # Agregar campos espec√≠ficos seg√∫n tipo
            if tipo_comision == "Encargados CM":
                reemplazos.update({
                    '<<CENTRO_MAESTROS>>': persona.get('centro_maestros', ''),
                    '<<DOMICILIO>>': persona.get('domicilio', ''),
                    '<<COLONIA>>': persona.get('colonia', ''),
                    '<<MUNICIPIO>>': persona.get('municipio', '')
                })
            else:  # Comisiones Generales
                reemplazos.update({
                    '<<INSTITUCION>>': persona.get('institucion', ''),
                    '<<UBICACION>>': persona.get('institucion', ''),  # Puede ser el mismo
                    '<<DOMICILIO>>': persona.get('domicilio', ''),
                    '<<COLONIA>>': persona.get('colonia', ''),
                    '<<MUNICIPIO>>': persona.get('municipio', ''),
                    '<<CP>>': persona.get('cp', '')
                })
            
            # Reemplazar en p√°rrafos
            for paragraph in doc.paragraphs:
                texto_completo = paragraph.text
                for marcador, valor in reemplazos.items():
                    if marcador in texto_completo:
                        texto_completo = texto_completo.replace(marcador, str(valor))
                
                if texto_completo != paragraph.text:
                    for run in paragraph.runs:
                        run.text = ''
                    if paragraph.runs:
                        paragraph.runs[0].text = texto_completo
            
            # Reemplazar en tablas si existen
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            texto_completo = paragraph.text
                            for marcador, valor in reemplazos.items():
                                if marcador in texto_completo:
                                    texto_completo = texto_completo.replace(marcador, str(valor))
                            
                            if texto_completo != paragraph.text:
                                for run in paragraph.runs:
                                    run.text = ''
                                if paragraph.runs:
                                    paragraph.runs[0].text = texto_completo
            
            docs.append(doc)
            oficio_actual += 1
        
        # Combinar documentos SIN docxcompose
        doc_final = docs[0]
        
        for doc in docs[1:]:
            # Copiar TODO excepto sectPr
            for element in list(doc.element.body):
                if element.tag.endswith('sectPr'):
                    continue
                doc_final.element.body.append(element)
        
        # Guardar
        tipo_archivo = "Encargados_CM" if tipo_comision == "Encargados CM" else "Comisiones_Generales"
        output_path = os.path.join(os.path.dirname(__file__), f'{tipo_archivo}_{oficio_inicial}.docx')
        doc_final.save(output_path)
        
        return output_path
        
    except Exception as e:
        import traceback
        error_completo = traceback.format_exc()
        raise Exception(f"Error al generar comisiones: {str(e)}\n\nStack trace:\n{error_completo}")

# ============= LOGIN =============
if 'logged_in' not in st.session_state:
    st.session_state['logged_in'] = False

if not st.session_state['logged_in']:
    st.title("üîê Sistema de Gesti√≥n de RH DFC")
    st.markdown("**Direcci√≥n de Formaci√≥n Continua** - Secretar√≠a de Educaci√≥n Jalisco")
    st.markdown("---")
    
    col1, col2, col3 = st.columns([1,2,1])
    with col2:
        st.subheader("Iniciar Sesi√≥n")
        usuario = st.text_input("Usuario")
        password = st.text_input("Contrase√±a", type="password")
        
        if st.button("Ingresar", use_container_width=True, type="primary"):
            valido, nombre, tipo = verificar_login(usuario, password)
            if valido:
                st.session_state['logged_in'] = True
                st.session_state['usuario'] = usuario
                st.session_state['nombre_usuario'] = nombre
                st.session_state['tipo_usuario'] = tipo
                
                # Verificar alertas cr√≠ticas
                fechas_limite_login = verificar_fechas_limite()
                
                if fechas_limite_login['criticas']:
                    st.session_state['mostrar_alerta_login'] = True
                    st.session_state['alertas_criticas'] = fechas_limite_login['criticas']
                
                st.rerun()  
            else:
                st.error("‚ùå Usuario o contrase√±a incorrectos")
    st.stop()

# ============= VISORES (SOLO LECTURA) =============
tipo_usuario = st.session_state.get('tipo_usuario', 'admin')

if tipo_usuario == 'visor_viaticos':
    st.title("üìã Consulta de Empleados - Vi√°ticos üöó")
    st.markdown("**Vista de solo lectura**")
    
    col1, col2 = st.columns([4,1])
    with col2:
        st.write(f"üë§ **{st.session_state['nombre_usuario']}**")
        if st.button("üö™ Cerrar Sesi√≥n"):
            st.session_state['logged_in'] = False
            st.rerun()
    
    st.markdown("---")
    
    # Cargar datos
    if 'df_empleados' not in st.session_state:
        client = conectar_sheets()
        if client:
            spreadsheet = client.open("Dias_Economicos_Formacion_Continua")
            st.session_state['df_empleados'] = pd.DataFrame(spreadsheet.worksheet("Empleados").get_all_records())
    
    df_empleados = st.session_state['df_empleados'].copy()
    
    # B√∫squeda
    busqueda = st.text_input("üîç Buscar por nombre, RFC o CURP")
    
    if busqueda:
        mascara = (
            df_empleados['PATERNO'].str.contains(busqueda, case=False, na=False) |
            df_empleados['MATERNO'].str.contains(busqueda, case=False, na=False) |
            df_empleados['NOMBRE'].str.contains(busqueda, case=False, na=False) |
            df_empleados['RFC'].str.contains(busqueda, case=False, na=False) |
            df_empleados['CURP'].str.contains(busqueda, case=False, na=False)
        )
        df_empleados = df_empleados[mascara]
    
    st.info(f"üìä Mostrando {len(df_empleados)} empleados")
    
    # Vista restringida: solo CURP y RFC
    df_vista = df_empleados[['PATERNO', 'MATERNO', 'NOMBRE', 'CURP', 'RFC']].copy()
    df_vista['NOMBRE COMPLETO'] = df_vista['PATERNO'] + ' ' + df_vista['MATERNO'] + ' ' + df_vista['NOMBRE']
    df_vista = df_vista[['NOMBRE COMPLETO', 'CURP', 'RFC']]
    
    st.dataframe(df_vista, use_container_width=True, hide_index=True)
    st.stop()

elif tipo_usuario == 'visor_secretarias':
    st.title("üìã Directorio de Empleados - Secretarias")
    st.markdown("**Vista de solo lectura**")
    
    col1, col2 = st.columns([4,1])
    with col2:
        st.write(f"üë§ **{st.session_state['nombre_usuario']}**")
        if st.button("üö™ Cerrar Sesi√≥n"):
            st.session_state['logged_in'] = False
            st.rerun()
    
    st.markdown("---")
    
    # Cargar datos
    if 'df_empleados' not in st.session_state:
        client = conectar_sheets()
        if client:
            spreadsheet = client.open("Dias_Economicos_Formacion_Continua")
            st.session_state['df_empleados'] = pd.DataFrame(spreadsheet.worksheet("Empleados").get_all_records())
    
    df_empleados = st.session_state['df_empleados'].copy()
    
    # B√∫squeda
    busqueda = st.text_input("üîç Buscar por nombre, RFC, CURP o Centro de Maestros")
    
    if busqueda:
        mascara = (
            df_empleados['PATERNO'].str.contains(busqueda, case=False, na=False) |
            df_empleados['MATERNO'].str.contains(busqueda, case=False, na=False) |
            df_empleados['NOMBRE'].str.contains(busqueda, case=False, na=False) |
            df_empleados['RFC'].str.contains(busqueda, case=False, na=False) |
            df_empleados['CURP'].str.contains(busqueda, case=False, na=False) |
            df_empleados.get('CENTRO DE TRABAJO', pd.Series()).str.contains(busqueda, case=False, na=False)
        )
        df_empleados = df_empleados[mascara]
    
    st.info(f"üìä Mostrando {len(df_empleados)} empleados")
    
    # Vista: CURP, RFC, Tel√©fono, Centro de Maestros
    columnas = ['PATERNO', 'MATERNO', 'NOMBRE', 'CURP', 'RFC']
    
    # Agregar tel√©fono si existe
    if 'TELEFONO' in df_empleados.columns:
        columnas.append('TELEFONO')
    
    # Agregar centro de trabajo
    if 'CENTRO DE TRABAJO' in df_empleados.columns:
        columnas.append('CENTRO DE TRABAJO')
    
    df_vista = df_empleados[columnas].copy()
    df_vista['NOMBRE COMPLETO'] = df_vista['PATERNO'] + ' ' + df_vista['MATERNO'] + ' ' + df_vista['NOMBRE']
    
    cols_finales = ['NOMBRE COMPLETO', 'CURP', 'RFC']
    if 'TELEFONO' in columnas:
        cols_finales.append('TELEFONO')
    if 'CENTRO DE TRABAJO' in columnas:
        cols_finales.append('CENTRO DE TRABAJO')
    
    df_vista = df_vista[cols_finales]
    
    st.dataframe(df_vista, use_container_width=True, hide_index=True)
    st.stop()

# Si es admin, contin√∫a con la app normal
# ============= MAIN APP =============
st.title("üìÖ Sistema de Gesti√≥n de RH DFC")
st.markdown("**Direcci√≥n de Formaci√≥n Continua** - Secretar√≠a de Educaci√≥n Jalisco")

col1, col2 = st.columns([4,1])
with col2:
    st.write(f"üë§ **{st.session_state['nombre_usuario']}**")
    if st.button("üö™ Cerrar Sesi√≥n"):
        st.session_state['logged_in'] = False
        st.rerun()

st.markdown("---")

# CSS para mejorar visibilidad de pesta√±as
st.markdown("""
<style>
    /* Pesta√±as normales */
    .stTabs [data-baseweb="tab"] {
        height: 60px;
        padding: 15px 25px;
        background-color: #f0f2f6;
        border-radius: 5px 5px 0px 0px;
        margin-right: 5px;
        font-weight: 700;
        font-size: 18px !important;
        transition: all 0.3s ease;
    }
    
    /* Texto dentro de pesta√±as m√°s grande */
    .stTabs [data-baseweb="tab"] p {
        font-size: 18px !important;
        font-weight: 700 !important;
    }
    
    /* Hover sobre pesta√±as */
    .stTabs [data-baseweb="tab"]:hover {
        background-color: #e0e5eb;
        transform: translateY(-2px);
        box-shadow: 0 2px 5px rgba(0,0,0,0.1);
    }
    
    /* Pesta√±a activa/seleccionada */
    .stTabs [aria-selected="true"] {
        background-color: #0068c9 !important;
        color: white !important;
        font-weight: 800 !important;
        font-size: 19px !important;
        box-shadow: 0 4px 10px rgba(0,104,201,0.3);
    }
</style>
""", unsafe_allow_html=True)

# Cargar datos solo una vez al inicio

# Mostrar alerta de login si hay propuestas cr√≠ticas
if st.session_state.get('mostrar_alerta_login', False):
    with st.container():
        st.error("### ‚ö†Ô∏è RECUERDA FECHAS L√çMITE PARA CAPTURA EN RH ‚ö†Ô∏è")
        
        for item in st.session_state.get('alertas_criticas', []):
            st.warning(f"üö® {item['sistema']} {item['quincena']}: Faltan {item['dias']} d√≠as - L√≠mite: {item['fecha']}")
        
        col1, col2 = st.columns([3, 1])
        with col2:
            if st.button("‚úÖ Entendido", type="primary"):
                st.session_state['mostrar_alerta_login'] = False
                st.rerun()
        
        st.markdown("---")

# Cargar datos solo una vez al inicio
if 'df_empleados' not in st.session_state:
    client = conectar_sheets()
    if client:
        try:
            spreadsheet = client.open("Dias_Economicos_Formacion_Continua")
            
            # Leer TODAS las hojas y convertir a DataFrames
            st.session_state['df_empleados'] = pd.DataFrame(spreadsheet.worksheet("Empleados").get_all_records())
            
            
            st.session_state['df_solicitudes'] = pd.DataFrame(spreadsheet.worksheet("Solicitudes").get_all_records())
            
            
            # Incapacidades con columnas por defecto
            df_incap = pd.DataFrame(spreadsheet.worksheet("Incapacidades").get_all_records())
           
            if len(df_incap) == 0:
                df_incap = pd.DataFrame(columns=['ID', 'EmpleadoID', 'RFC', 'Nombre Completo', 'Correo Empleado',
                                                  'Telefono Contacto', 'Numero Incapacidad', 'Fecha Inicio', 
                                                  'Fecha Termino', 'Dias Totales', 'Tipo Incapacidad', 'Excede Dias',
                                                  'Dias Enfermedad General', 'Dias Maternidad', 'Dias Riesgo Trabajo',
                                                  'Dias Posible Riesgo', 'Mes Correspondiente', 'Estado', 'Registrado Por'])
            st.session_state['df_incapacidades'] = df_incap
            
            # Pendientes con columnas por defecto
            df_pend = pd.DataFrame(spreadsheet.worksheet("Pendientes_Empleado").get_all_records())
            
            if len(df_pend) == 0:
                df_pend = pd.DataFrame(columns=['ID', 'EmpleadoID', 'RFC', 'Nombre Completo', 'Tipo_Pendiente',
                                                'Descripcion', 'Quincena', 'A√±o', 'Estado', 'Fecha_Registro',
                                                'Fecha_Completado', 'Completado_Por'])
            st.session_state['df_pendientes'] = df_pend
            
            st.session_state['df_constancias'] = pd.DataFrame(spreadsheet.worksheet("Constancias").get_all_records())
            # AGREGAR ESTA L√çNEA:
            st.session_state['df_comisiones'] = pd.DataFrame(spreadsheet.worksheet("Comisiones").get_all_records())
            # Guardar el cliente para escrituras
            st.session_state['client'] = client
            st.session_state['spreadsheet_name'] = "Dias_Economicos_Formacion_Continua"
            
        except Exception as e:
            st.error(f"Error al cargar datos: {str(e)}")
            st.stop()
    else:
        st.error("No se pudo conectar a Google Sheets")
        st.stop()

# Usar los DataFrames desde session_state
df_empleados = st.session_state['df_empleados'].copy()
df_solicitudes = st.session_state['df_solicitudes'].copy()
df_incapacidades = st.session_state['df_incapacidades'].copy()
df_pendientes = st.session_state['df_pendientes'].copy()
df_constancias = st.session_state['df_constancias'].copy()
df_comisiones = st.session_state['df_comisiones'].copy()

# Calcular d√≠as disponibles
for idx, emp in df_empleados.iterrows():
    emp_id = emp['ID']
    solicitudes_emp = df_solicitudes[df_solicitudes['EmpleadoID'] == emp_id]
    
    dias_usados = 0
    for _, sol in solicitudes_emp.iterrows():
        if sol['Tipo Permiso'] == 'economico':
            dias_usados += int(sol['Dias Solicitados'])
    
    df_empleados.at[idx, 'DIAS_REALES'] = 9 - dias_usados

# SIDEBAR: Alertas
with st.sidebar:
    st.header("üîî Alertas y Notificaciones")
    
    # Alertas de d√≠as disponibles
    if len(df_empleados) > 0:
        alertas = generar_alertas(df_empleados)
        
        if alertas:
            for alerta in alertas:
                if alerta['tipo'] == 'error':
                    st.error(alerta['mensaje'])
                elif alerta['tipo'] == 'warning':
                    st.warning(alerta['mensaje'])
                else:
                    st.info(alerta['mensaje'])
        else:
            st.success("‚úÖ No hay alertas de d√≠as")
    
    # Alertas de propuestas CR√çTICAS
    st.markdown("---")
    st.markdown("**üìã Propuestas Urgentes**")
    
    fechas_limite = verificar_fechas_limite()
    
    if fechas_limite['criticas']:
        for item in fechas_limite['criticas']:
            st.error(f"üö® {item['sistema']} {item['quincena']}: {item['dias']} d√≠as ({item['fecha']})")
    else:
        st.success("‚úÖ Sin propuestas urgentes")
    
    st.markdown("---")
    st.markdown("**üìä Resumen General**")
    if len(df_empleados) > 0:
        st.metric("Total Empleados", len(df_empleados))
        st.metric("Solicitudes Registradas", len(df_solicitudes))
        dias_promedio = df_empleados['DIAS_REALES'].mean()
        st.metric("D√≠as Disponibles (Promedio)", int(dias_promedio))

# TABS PRINCIPALES
tab1, tab2, tab3, tab4, tab5, tab6, tab7, tab8 = st.tabs([
    "üìù D√≠as Econ√≥micos",
    "üè• Incapacidades",
    "üë• Ver Empleados", 
    "üìä Estatus Individual",
    "üìÑ Reportes",
    "üîî Recordatorios",
    "üìã Gesti√≥n Documental",  # NUEVO
    "üìã Normativa"
])

# TAB 1: D√çAS ECON√ìMICOS
with tab1:
    st.header("Registrar Nueva Solicitud")
    
    if len(df_empleados) == 0:
        st.warning("‚ö†Ô∏è No hay empleados registrados")
    else:
        col1, col2 = st.columns(2)
        
        with col1:
            opciones = [(e['ID'], f"{e['PATERNO']} {e['MATERNO']} {e['NOMBRE']} - {e['PUESTO']} ({int(e['DIAS_REALES'])} d√≠as)") 
                        for _, e in df_empleados.iterrows()]
            emp_id = st.selectbox("Seleccionar Empleado", [o[0] for o in opciones], 
                                  format_func=lambda x: next(o[1] for o in opciones if o[0]==x))
            
            tipo = st.selectbox("Tipo de Permiso", list(NORMATIVA.keys()), 
                               format_func=lambda x: f"{NORMATIVA[x]['nombre']} (max. {NORMATIVA[x]['max_dias']} d√≠as)")
        
        with col2:
            dias = st.number_input("N√∫mero de D√≠as", 1, NORMATIVA[tipo]['max_dias'], 1)
            aprobado = st.text_input("Aprobado Por", "Jefe de Departamento")
        
        st.markdown("---")
        st.subheader("üìÖ Fechas Solicitadas")
        
        # Selector de tipo de fechas
        tipo_fechas = st.radio(
            "¬øC√≥mo quieres ingresar las fechas?",
            ["Consecutivas (rango)", "NO consecutivas (manual)"],
            horizontal=True
        )
        
        fechas_procesadas = []
        
        if tipo_fechas == "Consecutivas (rango)":
            col_f1, col_f2 = st.columns(2)
            with col_f1:
                fecha_inicio_input = st.date_input("Fecha de Inicio", value=datetime.now())
            with col_f2:
                fecha_fin_input = st.date_input("Fecha de Fin", value=datetime.now() + timedelta(days=dias-1))
            
            # Generar todas las fechas del rango
            fecha_actual = fecha_inicio_input
            while fecha_actual <= fecha_fin_input:
                fechas_procesadas.append(datetime.combine(fecha_actual, datetime.min.time()))
                fecha_actual += timedelta(days=1)
            
            dias_rango = len(fechas_procesadas)
            if dias_rango == dias:
                st.success(f"‚úÖ {dias_rango} fecha(s): {fecha_inicio_input.strftime('%d/%m/%Y')} al {fecha_fin_input.strftime('%d/%m/%Y')}")
            else:
                st.warning(f"‚ö†Ô∏è El rango tiene {dias_rango} d√≠as pero solicitaste {dias}")
        
        else:
            fechas_input = st.text_input(
                "Escribe las fechas separadas por comas (formato: dd/mm/yyyy)",
                placeholder="Ejemplo: 05/01/2025, 10/01/2025, 20/01/2025",
                help="Puedes solicitar d√≠as NO consecutivos"
            )
            
            if fechas_input:
                try:
                    for f in fechas_input.split(','):
                        fecha_obj = datetime.strptime(f.strip(), '%d/%m/%Y')
                        fechas_procesadas.append(fecha_obj)
                    fechas_procesadas.sort()
                    st.success(f"‚úÖ {len(fechas_procesadas)} fecha(s) v√°lida(s): {', '.join([f.strftime('%d/%m/%Y') for f in fechas_procesadas])}")
                    
                    if len(fechas_procesadas) != dias:
                        st.warning(f"‚ö†Ô∏è Solicitaste {dias} d√≠as pero ingresaste {len(fechas_procesadas)} fechas")
                except:
                    st.error("‚ùå Formato incorrecto. Usa: dd/mm/yyyy, dd/mm/yyyy")
        
        motivo = st.text_area("Motivo/Descripci√≥n", height=100)
        
        # Info empleado
        emp_info = df_empleados[df_empleados['ID']==emp_id].iloc[0]
        st.info(f"""
        **üìã Informaci√≥n del Empleado:**
        - **RFC:** {emp_info['RFC']}
        - **Puesto:** {emp_info['PUESTO']}
        - **Centro de Trabajo:** {emp_info.get('CENTRO DE TRABAJO', 'N/A')}
        - **D√≠as Disponibles:** **{emp_info['DIAS_REALES']}/9**
        """)
        # Verificar si hay concentraci√≥n de personal
        if fechas_procesadas and tipo == 'economico':
            empleados_ausentes = []
            
            for fecha_check in fechas_procesadas:
                # Contar cu√°ntos empleados estar√°n ausentes ese d√≠a
                for _, sol in df_solicitudes.iterrows():
                    sol_inicio = pd.to_datetime(sol['Fecha Inicio'])
                    sol_fin = pd.to_datetime(sol['Fecha Fin'])
                    
                    if sol_inicio.date() <= fecha_check.date() <= sol_fin.date():
                        empleados_ausentes.append({
                            'fecha': fecha_check.strftime('%d/%m/%Y'),
                            'nombre': sol['Nombre Completo'],
                            'tipo': sol['Tipo Permiso']
                        })
            
            # Contar por fecha
            from collections import Counter
            fechas_count = Counter([e['fecha'] for e in empleados_ausentes])
            
            # Alertar si alguna fecha tiene 5+
            for fecha, count in fechas_count.items():
                if count >= 4:  # 4 existentes + 1 nuevo = 5 total
                    st.error(f"""
                    üö® **ALERTA: CONCENTRACI√ìN DE PERSONAL**
                    
                    El **{fecha}** ya tienen permiso **{count} empleados**
                    
                    Si registras esta solicitud ser√°n **{count + 1} empleados ausentes**
                    
                    ‚ö†Ô∏è **IMPACTO OPERATIVO:** Posible desabasto de personal
                    """)
                    
                    # Mostrar qui√©nes estar√°n ausentes
                    ausentes_fecha = [e for e in empleados_ausentes if e['fecha'] == fecha]
                    for aus in ausentes_fecha[:5]:
                        st.warning(f"‚Ä¢ {aus['nombre']} - {aus['tipo']}")
                    
                    if len(ausentes_fecha) > 5:
                        st.warning(f"... y {len(ausentes_fecha) - 5} m√°s")
                
                st.markdown("---")
        
        if st.button("‚úÖ REGISTRAR SOLICITUD", type="primary", use_container_width=True):
            if not fechas_procesadas:
                st.error("‚ùå Debes ingresar al menos una fecha v√°lida")
            elif len(fechas_procesadas) != dias:
                st.error(f"‚ùå El n√∫mero de fechas ({len(fechas_procesadas)}) no coincide con los d√≠as solicitados ({dias})")
            else:
                fecha_inicio = fechas_procesadas[0]
                fecha_fin = fechas_procesadas[-1]
                
                errores, advertencias = validar_solicitud(emp_id, tipo, dias, fecha_inicio, df_empleados, df_solicitudes)
                
                if errores:
                    st.error("**‚ùå SOLICITUD RECHAZADA**")
                    for error in errores:
                        st.error(error)
                else:
                    for adv in advertencias:
                        st.warning(adv)
                    
                    nombre = f"{emp_info['PATERNO']} {emp_info['MATERNO']} {emp_info['NOMBRE']}"
                    fechas_str = ", ".join([f.strftime('%d/%m/%Y') for f in fechas_procesadas])
                    
                    nuevo_id = len(df_solicitudes) + 1
                    nueva_fila = [
                        nuevo_id, emp_id, emp_info['RFC'], nombre, tipo,
                        fecha_inicio.strftime('%Y-%m-%d'),
                        fecha_fin.strftime('%Y-%m-%d'),
                        dias,
                        f"{motivo} | Fechas: {fechas_str}",
                        datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                        aprobado,
                        st.session_state['nombre_usuario']
                    ]
                    
                    # RECONECTAR para escribir
                    client = st.session_state['client']
                    spreadsheet = client.open(st.session_state['spreadsheet_name'])
                    sheet_sol = spreadsheet.worksheet("Solicitudes")
                    sheet_emp = spreadsheet.worksheet("Empleados")

                    # ESCRIBIR
                    sheet_sol.append_row(nueva_fila)

                    # Actualizar session_state
                    st.session_state['df_solicitudes'] = pd.DataFrame(sheet_sol.get_all_records())
                    st.session_state['df_empleados'] = pd.DataFrame(sheet_emp.get_all_records())
                    dias_restantes = int(emp_info['DIAS_REALES'] - dias) if tipo == 'economico' else int(emp_info['DIAS_REALES'])
                    
                    # CONFIRMACI√ìN
                    st.success("# ‚úÖ ¬°SOLICITUD REGISTRADA EXITOSAMENTE!")
                    st.balloons()
                    st.success(f"### üìã Folio: {nuevo_id}")
                    st.success(f"### üë§ {nombre}")
                    st.success(f"### üìÖ Fechas: {fechas_str}")
                    st.success(f"### üìä D√≠as restantes: **{dias_restantes}/9**")
                    st.success(f"### ‚úçÔ∏è Registrado por: {st.session_state['nombre_usuario']}")
                    st.toast(f"‚úÖ Solicitud #{nuevo_id} registrada", icon="‚úÖ")
                    
                    if st.button("üîÑ Registrar Otra Solicitud"):
                        st.rerun()

# TAB 2: INCAPACIDADES
with tab2:
    st.header("üè• Registro de Incapacidades")
    
    if len(df_empleados) == 0:
        st.warning("‚ö†Ô∏è No hay empleados registrados")
    else:
        col1, col2 = st.columns(2)
        
        with col1:
            opciones_inc = [(e['ID'], f"{e['PATERNO']} {e['MATERNO']} {e['NOMBRE']} - {e['RFC']}") 
                           for _, e in df_empleados.iterrows()]
            emp_id_inc = st.selectbox("Seleccionar Empleado", [o[0] for o in opciones_inc], 
                                      format_func=lambda x: next(o[1] for o in opciones_inc if o[0]==x), key="emp_incap")
            
            num_incapacidad = st.text_input("N√∫mero de Incapacidad (Folio IMSS)", placeholder="123456789")
            
            tipo_incap = st.selectbox("Tipo de Incapacidad", [
                "Enfermedad General",
                "Maternidad",
                "Riesgo de Trabajo",
                "Posible Riesgo de Trabajo"
            ])
        
        with col2:
            fecha_inicio_inc = st.date_input("Fecha Inicio", value=datetime.now(), key="fecha_inicio_inc")
            fecha_termino_inc = st.date_input("Fecha T√©rmino", value=datetime.now() + timedelta(days=3), key="fecha_termino_inc")
            
            dias_totales = (fecha_termino_inc - fecha_inicio_inc).days + 1
            st.metric("D√≠as Totales", dias_totales)
            
            correo_emp = st.text_input("Correo Empleado", placeholder="empleado@educacion.gob.mx")
            telefono_emp = st.text_input("Tel√©fono Contacto", placeholder="3312345678")
        
        # Info empleado
        emp_info_inc = df_empleados[df_empleados['ID']==emp_id_inc].iloc[0]
        
        # Calcular d√≠as acumulados en el a√±o
        incap_empleado = df_incapacidades[df_incapacidades['EmpleadoID'] == emp_id_inc]
        if len(incap_empleado) > 0:
            incap_empleado['Fecha_Inicio'] = pd.to_datetime(incap_empleado['Fecha Inicio'], errors='coerce')
            incap_a√±o = incap_empleado[incap_empleado['Fecha_Inicio'].dt.year == datetime.now().year]
            dias_acumulados = int(incap_a√±o['Dias Totales'].sum())
        else:
            dias_acumulados = 0
        
        dias_con_nueva = dias_acumulados + dias_totales
        excede = dias_con_nueva > 28  # L√≠mite com√∫n Art. 44
        
        if excede:
            st.error(f"""
            ‚ö†Ô∏è **ALERTA ART√çCULO 44**
            
            Este empleado acumular√≠a **{dias_con_nueva} d√≠as** de incapacidad en {datetime.now().year}
            (Actual: {dias_acumulados} + Nueva: {dias_totales})
            
            **EXCEDE EL L√çMITE DE 28 D√çAS**
            
            Acciones requeridas:
            - ‚úì Anexar Acta Circunstanciada
            - ‚úì Anexar Oficio
            - ‚úì Aplicar Art√≠culo 44
            """)
        elif dias_con_nueva > 20:
            st.warning(f"‚ö†Ô∏è Precauci√≥n: Acumular√≠a {dias_con_nueva} d√≠as de incapacidad en el a√±o (l√≠mite: 28)")
        
        st.info(f"""
        **üìã Informaci√≥n del Empleado:**
        - **RFC:** {emp_info_inc['RFC']}
        - **Puesto:** {emp_info_inc['PUESTO']}
        - **D√≠as acumulados en {datetime.now().year}:** {dias_acumulados}
        """)
        
        st.markdown("---")
        
        # Verificar concentraci√≥n de ausencias
        if st.button("‚úÖ REGISTRAR INCAPACIDAD", type="primary", use_container_width=True, key="btn_incap"):
            
            # Verificar concentraci√≥n de ausencias en esas fechas
            ausentes = []
            
            # Revisar solicitudes
            for _, sol in df_solicitudes.iterrows():
                sol_inicio = pd.to_datetime(sol['Fecha Inicio'])
                sol_fin = pd.to_datetime(sol['Fecha Fin'])
                
                if (fecha_inicio_inc <= sol_fin.date() and fecha_termino_inc >= sol_inicio.date()):
                    ausentes.append({
                        'nombre': sol['Nombre Completo'],
                        'tipo': sol['Tipo Permiso'],
                        'inicio': sol_inicio.strftime('%d/%m/%Y'),
                        'fin': sol_fin.strftime('%d/%m/%Y')
                    })
            
            # Revisar incapacidades
            for _, inc in df_incapacidades.iterrows():
                inc_inicio = pd.to_datetime(inc['Fecha Inicio'], errors='coerce')
                inc_fin = pd.to_datetime(inc['Fecha Termino'], errors='coerce')
                
                if pd.notna(inc_inicio) and pd.notna(inc_fin):
                    if (fecha_inicio_inc <= inc_fin.date() and fecha_termino_inc >= inc_inicio.date()):
                        ausentes.append({
                            'nombre': inc['Nombre Completo'],
                            'tipo': 'Incapacidad',
                            'inicio': inc_inicio.strftime('%d/%m/%Y'),
                            'fin': inc_fin.strftime('%d/%m/%Y')
                        })
            
            # Alerta si hay 5 o m√°s ausentes
            if len(ausentes) >= 5:
                st.warning(f"""
                ‚ö†Ô∏è **ALERTA DE CONCENTRACI√ìN DE PERSONAL**
                
                Del {fecha_inicio_inc.strftime('%d/%m/%Y')} al {fecha_termino_inc.strftime('%d/%m/%Y')}:
                
                **{len(ausentes)} empleados ausentes simult√°neamente:**
                """)
                for aus in ausentes[:10]:  # Mostrar m√°ximo 10
                    st.warning(f"‚Ä¢ {aus['nombre']} - {aus['tipo']} ({aus['inicio']} - {aus['fin']})")
                
                if len(ausentes) > 10:
                    st.warning(f"... y {len(ausentes) - 10} m√°s")
                
                st.warning("‚ö†Ô∏è **IMPACTO OPERATIVO:** Posible desabasto de personal")
            
            # Registrar incapacidad
            nombre = f"{emp_info_inc['PATERNO']} {emp_info_inc['MATERNO']} {emp_info_inc['NOMBRE']}"
            mes_corresp = fecha_inicio_inc.strftime('%B %Y')
            
            # Calcular d√≠as por tipo
            dias_por_tipo = {'Enfermedad General': 0, 'Maternidad': 0, 'Riesgo de Trabajo': 0, 'Posible Riesgo de Trabajo': 0}
            dias_por_tipo[tipo_incap] = dias_totales
            
            nueva_incap = [
                len(df_incapacidades) + 1,
                emp_id_inc,
                emp_info_inc['RFC'],
                nombre,
                correo_emp,
                telefono_emp,
                num_incapacidad,
                fecha_inicio_inc.strftime('%Y-%m-%d'),
                fecha_termino_inc.strftime('%Y-%m-%d'),
                dias_totales,
                tipo_incap,
                'S√ç' if excede else 'NO',
                dias_por_tipo['Enfermedad General'],
                dias_por_tipo['Maternidad'],
                dias_por_tipo['Riesgo de Trabajo'],
                dias_por_tipo['Posible Riesgo de Trabajo'],
                mes_corresp,
                'Pendiente',
                st.session_state['nombre_usuario']
            ]
            
            # RECONECTAR para escribir
            client = st.session_state['client']
            spreadsheet = client.open(st.session_state['spreadsheet_name'])
            sheet_incap = spreadsheet.worksheet("Incapacidades")

            # ESCRIBIR
            sheet_incap.append_row(nueva_incap)

            # Actualizar session_state
            st.session_state['df_incapacidades'] = pd.DataFrame(sheet_incap.get_all_records())
                        
            st.success("# ‚úÖ ¬°INCAPACIDAD REGISTRADA!")
            st.balloons()
            st.success(f"### üìã Folio: {len(df_incapacidades) + 1}")
            st.success(f"### üë§ {nombre}")
            st.success(f"### üìÖ Del {fecha_inicio_inc.strftime('%d/%m/%Y')} al {fecha_termino_inc.strftime('%d/%m/%Y')}")
            st.success(f"### üïí D√≠as: **{dias_totales}**")
            st.success(f"### üìä Total acumulado {datetime.now().year}: **{dias_con_nueva} d√≠as**")
            if excede:
                st.error("### ‚ö†Ô∏è REQUIERE APLICAR ART√çCULO 44")
            st.toast(f"‚úÖ Incapacidad registrada", icon="‚úÖ")
            
            if st.button("üîÑ Registrar Otra Incapacidad"):
                st.rerun()
    
    # Mostrar incapacidades recientes
    st.markdown("---")
    st.subheader("üìã Incapacidades Recientes")
    if len(df_incapacidades) > 0:
        columnas_mostrar = ['Nombre Completo', 'Numero Incapacidad', 'Fecha Inicio', 'Fecha Termino', 
                           'Dias Totales', 'Tipo Incapacidad', 'Excede Dias', 'Estado']
        st.dataframe(df_incapacidades[columnas_mostrar].tail(10), use_container_width=True, hide_index=True)
    else:
        st.info("No hay incapacidades registradas")

# TAB 3: VER EMPLEADOS
with tab3:
    st.header("üë• Plantilla de Personal")
    
    if len(df_empleados) > 0:
        col1, col2 = st.columns([2,1])
        with col1:
            busqueda = st.text_input("üîç Buscar por nombre, RFC o puesto")
        with col2:
            if st.button("üîÑ Actualizar Datos"):
                st.rerun()
        
        df_filtrado = df_empleados.copy()
        if busqueda:
            mascara = (
                df_filtrado['PATERNO'].str.contains(busqueda, case=False, na=False) |
                df_filtrado['MATERNO'].str.contains(busqueda, case=False, na=False) |
                df_filtrado['NOMBRE'].str.contains(busqueda, case=False, na=False) |
                df_filtrado['RFC'].str.contains(busqueda, case=False, na=False) |
                df_filtrado['PUESTO'].str.contains(busqueda, case=False, na=False)
            )
            df_filtrado = df_filtrado[mascara]
        
        st.info(f"üìä Mostrando {len(df_filtrado)} de {len(df_empleados)} empleados")
        
        # Agregar columna de pendientes
        df_mostrar = df_filtrado.copy()
        pendientes_count = []
        for _, emp in df_mostrar.iterrows():
            pends = df_pendientes[
                (df_pendientes['EmpleadoID'] == emp['ID']) & 
                (df_pendientes['Estado'] == 'Pendiente')
            ]
            count = len(pends)
            
            if count > 0:
                descripciones = pends['Tipo_Pendiente'].head(2).tolist()
                texto = f"‚ö†Ô∏è {count}: {', '.join(descripciones)}"
                if count > 2:
                    texto += "..."
                pendientes_count.append(texto)
            else:
                pendientes_count.append('‚úÖ 0')
        
        df_mostrar['PENDIENTES'] = pendientes_count
        
        columnas_mostrar = ['RFC', 'PATERNO', 'MATERNO', 'NOMBRE', 'PUESTO', 'DIAS_REALES', 'PENDIENTES']
        df_display = df_mostrar[columnas_mostrar].copy()
        df_display = df_display.rename(columns={'DIAS_REALES': 'DIAS DISPONIBLES'})
        
        st.dataframe(df_display, use_container_width=True, hide_index=True)
    else:
        st.warning("No hay empleados registrados")

# TAB 4: ESTATUS INDIVIDUAL
with tab4:
    st.header("üìä Estatus Individual de Empleados")
    
    if len(df_empleados) > 0:
        busqueda = st.text_input("üîç Buscar empleado", key="busq_individual")
        
        df_filtrado = df_empleados
        if busqueda:
            mascara = (
                df_filtrado['PATERNO'].str.contains(busqueda, case=False, na=False) |
                df_filtrado['MATERNO'].str.contains(busqueda, case=False, na=False) |
                df_filtrado['NOMBRE'].str.contains(busqueda, case=False, na=False) |
                df_filtrado['RFC'].str.contains(busqueda, case=False, na=False)
            )
            df_filtrado = df_filtrado[mascara]
        
        for _, emp in df_filtrado.iterrows():
            nombre = f"{emp['PATERNO']} {emp['MATERNO']} {emp['NOMBRE']}"
            
            # Contar pendientes
            pendientes_emp = df_pendientes[
                (df_pendientes['EmpleadoID'] == emp['ID']) & 
                (df_pendientes['Estado'] == 'Pendiente')
            ]
            num_pendientes = len(pendientes_emp)
            
            titulo = f"üë§ {nombre} - {emp['PUESTO']}"
            if num_pendientes > 0:
                titulo += f" ‚ö†Ô∏è {num_pendientes} pendiente(s)"
            
            with st.expander(titulo):
                col1, col2, col3, col4 = st.columns(4)
                
                dias_disp = int(emp['DIAS_REALES'])
                color = "üü¢" if dias_disp > 3 else "üü°" if dias_disp > 1 else "üî¥"
                
                with col1:
                    st.metric("D√≠as Disponibles", f"{color} {dias_disp}/9")
                with col2:
                    solicitudes_emp = df_solicitudes[df_solicitudes['EmpleadoID'] == emp['ID']]
                    st.metric("Total Solicitudes", len(solicitudes_emp))
                with col3:
                    st.metric("RFC", emp['RFC'])
                with col4:
                    st.metric("Pendientes", f"{'‚ö†Ô∏è' if num_pendientes > 0 else '‚úÖ'} {num_pendientes}")
                
                st.markdown("**Informaci√≥n Completa:**")
                info_cols = st.columns(2)
                with info_cols[0]:
                    st.write(f"**CURP:** {emp.get('CURP', 'N/A')}")
                    st.write(f"**Plaza:** {emp.get('PLAZA', 'N/A')}")
                with info_cols[1]:
                    st.write(f"**Centro:** {emp.get('CENTRO DE TRABAJO', 'N/A')}")
                    st.write(f"**Quincena:** {emp.get('QNA FIN', 'N/A')}")
                
                # PENDIENTES
                st.markdown("---")
                if num_pendientes > 0:
                    st.markdown("### ‚ö†Ô∏è PENDIENTES ACTIVOS:")
                    for _, pend in pendientes_emp.iterrows():
                        col_p1, col_p2 = st.columns([3,1])
                        with col_p1:
                            st.error(f"""
                            **{pend['Tipo_Pendiente']}:** {pend['Descripcion']}  
                            Registrado: {pend['Fecha_Registro']}
                            """)
                        with col_p2:
                            if st.button("‚úÖ Completar", key=f"comp_{pend['ID']}"):
                                # RECONECTAR para escribir
                                client = st.session_state['client']
                                spreadsheet = client.open(st.session_state['spreadsheet_name'])
                                sheet_pend = spreadsheet.worksheet("Pendientes_Empleado")
                                
                                # Buscar en la COLUMNA A (ID) espec√≠ficamente
                                todos_ids = sheet_pend.col_values(1)  # Columna A = IDs
                                
                                try:
                                    fila = todos_ids.index(str(pend['ID'])) + 1  # +1 porque index empieza en 0
                                    
                                    # Actualizar las celdas
                                    sheet_pend.update_cell(fila, 9, 'Completado')
                                    sheet_pend.update_cell(fila, 11, datetime.now().strftime('%Y-%m-%d'))
                                    sheet_pend.update_cell(fila, 12, st.session_state['nombre_usuario'])
                                    
                                    # Actualizar session_state
                                    st.session_state['df_pendientes'] = pd.DataFrame(sheet_pend.get_all_records())
                                    
                                    st.success("‚úÖ Marcado como completado")
                                    st.rerun()
                                    
                                except ValueError:
                                    st.error(f"‚ùå No se encontr√≥ el pendiente ID {pend['ID']}")
                else:
                    st.success("### ‚úÖ SIN PENDIENTES - Todo al d√≠a")
                
                # Agregar nuevo pendiente
                with st.expander("‚ûï Agregar Nuevo Pendiente"):
                    tipo_pend = st.selectbox("Tipo", [
                        "N√≥mina (firma)",
                        "Constancia (entregar)",
                        "Comisi√≥n (recibir)",
                        "Posada (juguete/boleto)",
                        "Incapacidad (documentos)",
                        "Otro"
                    ], key=f"tipo_pend_{emp['ID']}")
                    
                    desc_pend = st.text_input("Descripci√≥n", 
                                              placeholder="Ej: Firma Quincena 02/2026", 
                                              key=f"desc_pend_{emp['ID']}")
                    
                    col_qna, col_a√±o = st.columns(2)
                    with col_qna:
                        qna_pend = st.text_input("Quincena (opcional)", 
                                                 placeholder="02", 
                                                 key=f"qna_pend_{emp['ID']}")
                    with col_a√±o:
                        a√±o_pend = st.number_input("A√±o", 
                                                   min_value=2024, 
                                                   max_value=2030, 
                                                   value=datetime.now().year,
                                                   key=f"a√±o_pend_{emp['ID']}")
                    
                    if st.button("Registrar Pendiente", key=f"reg_pend_{emp['ID']}"):
                        if desc_pend:
                            nuevo_pend = [
                                len(df_pendientes) + 1,
                                emp['ID'],
                                emp['RFC'],
                                nombre,
                                tipo_pend,
                                desc_pend,
                                qna_pend if qna_pend else '',
                                a√±o_pend,
                                'Pendiente',
                                datetime.now().strftime('%Y-%m-%d'),
                                '',
                                ''
                            ]
                            # RECONECTAR para escribir
                            client = st.session_state['client']
                            spreadsheet = client.open(st.session_state['spreadsheet_name'])
                            sheet_pend = spreadsheet.worksheet("Pendientes_Empleado")

                            # ESCRIBIR
                            sheet_pend.append_row(nuevo_pend)

                            # Actualizar session_state
                            st.session_state['df_pendientes'] = pd.DataFrame(sheet_pend.get_all_records())

                            st.success("‚úÖ Pendiente registrado")
                            st.rerun()
                        else:
                            st.error("La descripci√≥n es obligatoria")
                
                # Historial de solicitudes
                if len(solicitudes_emp) > 0:
                    st.markdown("---")
                    st.markdown("**üìã Historial de Solicitudes:**")
                    columnas = ['Tipo Permiso', 'Fecha Inicio', 'Fecha Fin', 'Dias Solicitados', 'Motivo', 'Aprobado Por']
                    if 'Registrado Por' in solicitudes_emp.columns:
                        columnas.append('Registrado Por')
                    st.dataframe(solicitudes_emp[columnas], use_container_width=True, hide_index=True)

# TAB 5: REPORTES
with tab5:
    st.header("üìÑ Generaci√≥n de Reportes")
    
    st.info("üéØ Reportes integrados con trazabilidad total y exportaci√≥n completa del mes")
    
    # Secci√≥n 1: REPORTE COMPLETO DEL MES
    st.markdown("### üì¶ Reporte Completo del Mes")
    st.markdown("Exporta **TODO** en un solo Excel: solicitudes, incapacidades, pendientes y estad√≠sticas")
    
    col_mes1, col_mes2, col_mes3 = st.columns(3)
    with col_mes1:
        mes_reporte = st.selectbox("Mes", range(1, 13), 
                                   index=datetime.now().month - 1,
                                   format_func=lambda x: datetime(2000, x, 1).strftime('%B'))
    with col_mes2:
        a√±o_reporte = st.number_input("A√±o", min_value=2024, max_value=2030, 
                                      value=datetime.now().year)
    with col_mes3:
        if st.button("üì• Generar Reporte Completo", use_container_width=True, type="primary"):
            with st.spinner("Generando reporte completo..."):
                excel_completo = generar_reporte_completo_mes(
                    df_empleados, df_solicitudes, df_incapacidades, df_pendientes,
                    mes_reporte, a√±o_reporte
                )
                
                mes_nombre = datetime(2000, mes_reporte, 1).strftime('%B')
                nombre_archivo = f"Reporte_Completo_{mes_nombre}_{a√±o_reporte}.xlsx"
                
                st.download_button(
                    "üíæ Descargar Reporte Completo",
                    excel_completo,
                    nombre_archivo,
                    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True
                )
                st.success(f"‚úÖ Reporte de {mes_nombre} {a√±o_reporte} generado")
    
    st.markdown("---")
    
    # Secci√≥n 2: REPORTES INDIVIDUALES
    st.markdown("### üìã Reportes Individuales")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.subheader("üì• Reportes de Empleados")
        if st.button("Descargar Plantilla (Excel)", use_container_width=True):
            if len(df_empleados) > 0:
                output = io.BytesIO()
                with pd.ExcelWriter(output, engine='openpyxl') as writer:
                    df_empleados.to_excel(writer, sheet_name='Empleados', index=False)
                
                st.download_button(
                    "üíæ Descargar Excel",
                    output.getvalue(),
                    f"empleados_{datetime.now().strftime('%Y%m%d')}.xlsx",
                    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                )
    
    with col2:
        st.subheader("üì• Reportes de Solicitudes")
        if st.button("Descargar Historial (Excel)", use_container_width=True):
            if len(df_solicitudes) > 0:
                output = io.BytesIO()
                with pd.ExcelWriter(output, engine='openpyxl') as writer:
                    df_solicitudes.to_excel(writer, sheet_name='Solicitudes', index=False)
                
                st.download_button(
                    "üíæ Descargar Excel",
                    output.getvalue(),
                    f"solicitudes_{datetime.now().strftime('%Y%m%d')}.xlsx",
                    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                )
    
    st.markdown("---")
    
    # Secci√≥n 3: TRAZABILIDAD COMPLETA
    st.markdown("### üîç Trazabilidad Total")
    st.markdown("Historial completo con todos los detalles: qui√©n registr√≥, qui√©n aprob√≥, cu√°ndo, d√≥nde")
    
    if st.button("üìä Ver Trazabilidad Completa", use_container_width=True):
        df_trazabilidad = crear_trazabilidad_completa(df_solicitudes, df_empleados)
        
        if len(df_trazabilidad) > 0:
            st.dataframe(df_trazabilidad, use_container_width=True, hide_index=True)
            
            # Bot√≥n de descarga
            output = io.BytesIO()
            with pd.ExcelWriter(output, engine='openpyxl') as writer:
                df_trazabilidad.to_excel(writer, sheet_name='Trazabilidad', index=False)
            
            st.download_button(
                "üíæ Descargar Trazabilidad",
                output.getvalue(),
                f"trazabilidad_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx",
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True
            )
        else:
            st.info("No hay datos para mostrar")
    
    st.markdown("---")
    
    # Secci√≥n 4: ESTAD√çSTICAS GENERALES
    st.markdown("### üìä Estad√≠sticas Generales")
    
    if len(df_empleados) > 0:
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.metric("Total Empleados", len(df_empleados))
        with col2:
            criticos = len(df_empleados[df_empleados['DIAS_REALES'] <= 1])
            st.metric("‚ö†Ô∏è Cr√≠ticos", criticos)
        with col3:
            total_dias = df_empleados['DIAS_REALES'].sum()
            st.metric("Total D√≠as Disponibles", int(total_dias))
        with col4:
            if len(df_solicitudes) > 0:
                df_solicitudes['Fecha_Reg'] = pd.to_datetime(df_solicitudes['Fecha Registro'], errors='coerce')
                dias_usados = df_solicitudes[df_solicitudes['Fecha_Reg'].dt.year == datetime.now().year]['Dias Solicitados'].sum()
                st.metric("D√≠as Usados (2025)", int(dias_usados))

# TAB 6: RECORDATORIOS
with tab6:
    st.header("üîî Recordatorios de Fechas L√≠mite")
    
    st.info("üìÖ Fechas l√≠mite para entregar propuestas de pago a RH Central")
    
    fechas_limite_tab = verificar_fechas_limite()
    
    # CR√çTICAS
    if fechas_limite_tab['criticas']:
        st.markdown("### üö® CR√çTICAS (0-3 d√≠as)")
        for item in fechas_limite_tab['criticas']:
            col1, col2, col3, col4 = st.columns([2, 1, 2, 1])
            with col1:
                st.error(f"**{item['sistema']} {item['quincena']}**")
            with col2:
                st.metric("D√≠as", item['dias'])
            with col3:
                st.error(f"L√≠mite: {item['fecha']}")
            with col4:
                st.error("‚ö†Ô∏è URGENTE")
        st.markdown("---")
    
    # PR√ìXIMAS
    if fechas_limite_tab['proximas']:
        st.markdown("### ‚ö†Ô∏è PR√ìXIMAS (4-15 d√≠as)")
        for item in fechas_limite_tab['proximas']:
            col1, col2, col3 = st.columns([2, 1, 2])
            with col1:
                st.warning(f"**{item['sistema']} {item['quincena']}**")
            with col2:
                st.metric("D√≠as", item['dias'])
            with col3:
                st.warning(f"L√≠mite: {item['fecha']}")
        st.markdown("---")
    
    # FUTURAS
    if fechas_limite_tab['futuras']:
        st.markdown("### ‚ÑπÔ∏è FUTURAS (16-90 d√≠as)")
        for item in fechas_limite_tab['futuras']:
            col1, col2, col3 = st.columns([2, 1, 2])
            with col1:
                st.info(f"**{item['sistema']} {item['quincena']}**")
            with col2:
                st.metric("D√≠as", item['dias'])
            with col3:
                st.info(f"L√≠mite: {item['fecha']}")

# TAB 7: GESTI√ìN DOCUMENTAL
with tab7:
    st.header("üìã Gesti√≥n Documental")
    
    tipo_doc = st.selectbox(
    "Tipo de documento",
    ["üìÑ Constancias", "üöó Comisiones", "üìã Propuestas/Oficios (pr√≥ximamente)"],
    help="Selecciona el tipo de documento a generar"
    )
    
    if tipo_doc == "üìÑ Constancias":
        st.markdown("---")
        st.subheader("Generador de Constancias de Servicio")
        
        # Usar datos ya cargados desde session_state
        df_constancias = st.session_state['df_constancias']
        
        if len(df_constancias) == 0:
            st.error("‚ùå No hay datos de empleados en la hoja Constancias")
            st.stop()
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            num_quincena = st.number_input("N√∫mero de Quincena", min_value=1, max_value=24, value=24)
        
        with col2:
            a√±o_const = st.number_input("A√±o", min_value=2024, max_value=2030, value=2025)
        
        with col3:
            fecha_const = st.date_input("Fecha de elaboraci√≥n", value=datetime.now())
        
        st.markdown("---")
        st.markdown("**Seleccionar empleados para generar constancias:**")
        
        # Lista de empleados
        lista_empleados = df_constancias['Nombre Completo'].tolist()
        
        empleados_seleccionados = st.multiselect(
            "Empleados",
            options=lista_empleados,
            default=lista_empleados,
            help="Por defecto est√°n todos seleccionados. Puedes deseleccionar los que no necesites."
        )
        
        st.info(f"üìä **{len(empleados_seleccionados)} empleados seleccionados** de {len(lista_empleados)} totales")
        
        if st.button("‚úÖ Generar Constancias", type="primary", use_container_width=True):
            if not empleados_seleccionados:
                st.error("‚ùå Debes seleccionar al menos un empleado")
            else:
                with st.spinner("Generando constancias..."):
                    try:
                        # Filtrar df_constancias solo con empleados seleccionados
                        df_filtrado = df_constancias[df_constancias['Nombre Completo'].isin(empleados_seleccionados)].copy()
                        
                        # Generar documento UNA SOLA VEZ con todos los datos
                        output_path = generar_constancias_word(
                            df_filtrado,
                            empleados_seleccionados,
                            num_quincena,
                            a√±o_const,
                            fecha_const
                        )
                        
                        st.success(f"‚úÖ Constancias generadas exitosamente para {len(empleados_seleccionados)} empleados")

                        col1, col2 = st.columns(2)

                        # Bot√≥n Word
                        with col1:
                            with open(output_path, 'rb') as f:
                                st.download_button(
                                    label="üìÑ Descargar Word",
                                    data=f,
                                    file_name=f"Constancias_Q{num_quincena}_{a√±o_const}.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    use_container_width=True
                                )

                        # Bot√≥n PDF
                        with col2:
                            with st.spinner("Convirtiendo a PDF..."):
                                pdf_path = convertir_word_a_pdf(output_path)
                                
                                if pdf_path and os.path.exists(pdf_path):
                                    with open(pdf_path, 'rb') as f:
                                        st.download_button(
                                            label="üìï Descargar PDF",
                                            data=f,
                                            file_name=f"Constancias_Q{num_quincena}_{a√±o_const}.pdf",
                                            mime="application/pdf",
                                            use_container_width=True
                                        )
                                else:
                                    st.warning("‚ö†Ô∏è Conversi√≥n a PDF no disponible en este sistema")
                    
                    except Exception as e:
                        st.error(f"‚ùå Error al generar constancias: {str(e)}")
                        st.error("Verifica que la plantilla Word est√© en la ubicaci√≥n correcta")
    
    elif tipo_doc == "üöó Comisiones":
        st.markdown("---")
        st.subheader("Generador de Comisiones")
        
        # Selector de tipo de comisi√≥n
        tipo_comision = st.radio(
            "Tipo de comisi√≥n",
            ["Encargados CM", "Comisiones Generales"],
            horizontal=True
        )
        
        st.markdown("---")
        
        # Cargar datos de comisiones desde Google Sheets
        if 'df_comisiones' not in st.session_state:
            st.error("‚ùå No hay hoja 'Comisiones' en Google Sheets")
            st.stop()
        
        df_comisiones_todas = st.session_state['df_comisiones']
        
        # Filtrar por tipo
        if tipo_comision == "Encargados CM":
            df_filtrado = df_comisiones_todas[df_comisiones_todas['tipo_comision'] == 'Encargado CM'].copy()
        else:
            df_filtrado = df_comisiones_todas[df_comisiones_todas['tipo_comision'] == 'General'].copy()
        
        if len(df_filtrado) == 0:
            st.warning(f"‚ö†Ô∏è No hay registros de tipo '{tipo_comision}'")
            st.stop()
        
        # Inputs
        col1, col2, col3 = st.columns(3)
        
        with col1:
            oficio_inicial = st.number_input("N√∫mero de Oficio Inicial", min_value=1, max_value=999, value=118)
        
        with col2:
            fecha_doc = st.date_input("Fecha del Documento", value=datetime.now())
        
        with col3:
            st.write("")  # Espaciador
        
        col4, col5 = st.columns(2)
        
        with col4:
            fecha_inicio = st.date_input("Comisi√≥n del", value=datetime(2026, 1, 1))
        
        with col5:
            fecha_fin = st.date_input("Hasta el", value=datetime(2026, 2, 28))
        
        st.markdown("---")
        st.markdown("**Seleccionar personas para generar comisiones:**")
        
        # Lista de personas
        lista_personas = df_filtrado['nombre_completo'].tolist()
        
        personas_seleccionadas = st.multiselect(
            "Personas",
            options=lista_personas,
            default=lista_personas,
            help="Por defecto est√°n todas seleccionadas"
        )
        
        st.info(f"üìä **{len(personas_seleccionadas)} personas seleccionadas** de {len(lista_personas)} totales")
        
        # Vista previa
        if personas_seleccionadas:
            with st.expander("üëÅÔ∏è Vista previa de oficios"):
                preview_data = []
                oficio_temp = oficio_inicial
                for nombre in personas_seleccionadas:
                    preview_data.append({
                        'Oficio': f"{oficio_temp}/52/2026",
                        'Nombre': nombre
                    })
                    oficio_temp += 1
                st.dataframe(preview_data, use_container_width=True, hide_index=True)
        
        if st.button("‚úÖ Generar Comisiones", type="primary", use_container_width=True):
            if not personas_seleccionadas:
                st.error("‚ùå Debes seleccionar al menos una persona")
            else:
                with st.spinner("Generando comisiones..."):
                    try:
                        # Filtrar DataFrame
                        df_seleccionado = df_filtrado[df_filtrado['nombre_completo'].isin(personas_seleccionadas)].copy()
                        
                        # Generar documento
                        output_path = generar_comisiones_word(
                            df_seleccionado,
                            tipo_comision,
                            oficio_inicial,
                            fecha_doc,
                            fecha_inicio,
                            fecha_fin
                        )
                        
                        st.success(f"‚úÖ Comisiones generadas exitosamente para {len(personas_seleccionadas)} personas")
                        
                        col1, col2 = st.columns(2)
                        
                        # Bot√≥n Word
                        with col1:
                            with open(output_path, 'rb') as f:
                                tipo_archivo = "Encargados_CM" if tipo_comision == "Encargados CM" else "Comisiones_Generales"
                                st.download_button(
                                    label="üìÑ Descargar Word",
                                    data=f,
                                    file_name=f"{tipo_archivo}_Oficio_{oficio_inicial}.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    use_container_width=True
                                )
                        
                        # Bot√≥n PDF
                        with col2:
                            with st.spinner("Convirtiendo a PDF..."):
                                pdf_path = convertir_word_a_pdf(output_path)
                                
                                if pdf_path and os.path.exists(pdf_path):
                                    with open(pdf_path, 'rb') as f:
                                        st.download_button(
                                            label="üìï Descargar PDF",
                                            data=f,
                                            file_name=f"{tipo_archivo}_Oficio_{oficio_inicial}.pdf",
                                            mime="application/pdf",
                                            use_container_width=True
                                        )
                                else:
                                    st.warning("‚ö†Ô∏è Conversi√≥n a PDF no disponible")
                    
                    except Exception as e:
                        st.error(f"‚ùå Error al generar comisiones: {str(e)}")
    else:
        st.info("üöß Esta funcionalidad estar√° disponible pr√≥ximamente")


# TAB 8: NORMATIVA
with tab8:
    st.header("üìã Normativa Aplicable")
    
    st.info("""
    **Reglamento de las Condiciones Generales de Trabajo**  
    Secretar√≠a de Educaci√≥n del Estado de Jalisco
    """)
    
    st.markdown("### D√≠as Econ√≥micos (Asuntos Particulares)")
    st.markdown("""
    - ‚úÖ Hasta **3 d√≠as h√°biles** por ocasi√≥n
    - ‚úÖ M√°ximo **3 ocasiones** por a√±o calendario
    - ‚úÖ Intervalo m√≠nimo de **30 d√≠as** (desde el √∫ltimo d√≠a usado hasta el inicio del siguiente)
    - ‚úÖ Otorgados por el Jefe de Dependencia
    """)
    
    st.markdown("---")
    st.markdown("### Otras Licencias con Goce de Sueldo")
    
    tabla = pd.DataFrame([
        {
            'Motivo': v['nombre'], 
            'Duraci√≥n': f"{v['max_dias']} d√≠a(s)", 
            'L√≠mite': v['limite'],
            'Condiciones': v['descripcion']
        }
        for v in NORMATIVA.values()
    ])
    st.dataframe(tabla, use_container_width=True, hide_index=True)
    
    st.markdown("---")
    st.warning("""
    **‚ö†Ô∏è IMPORTANTE - L√≠mites de Uso:**
    
    - **Matrimonio**: Solo 1 vez EN LA VIDA
    - **Jubilaci√≥n**: Solo 1 vez EN LA VIDA (cuando se jubila)
    - **Examen Profesional**: M√°ximo 3 veces (licenciatura, maestr√≠a, doctorado)
    - **Mudanza**: M√°ximo 2 veces por a√±o
    - **Fallecimiento**: Sin l√≠mite (puede ocurrir varias veces)
    - **D√≠as Econ√≥micos**: 3 ocasiones por a√±o calendario
    """)